"""
WordImageProcessor
==================
Handles three modes for a .docx file:

  fix    – Fix every existing inline/anchored image:
             • Remove clipping line-spacing rules
             • Resize to fit usable page width × scale
             • Optionally centre the paragraph

  insert – Replace marker paragraphs with images.
             The marker is the image filename without extension
             (e.g. a paragraph containing only "diagram1" is replaced
              by the image file "diagram1.png").

  both   – fix then insert in one pass.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import List, Tuple, Dict, Any

from docx import Document
from docx.shared import Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
from PIL import Image

# EMU per centimetre
EMU_PER_CM = 360000
# Maximum sensible image width (15 inches)
MAX_WIDTH_EMU = 914400 * 15


# ──────────────────────────────────────────────────────────────
#  Helper: usable page width in EMU
# ──────────────────────────────────────────────────────────────

def _usable_width_emu(doc: Document, fallback_cm: float = 16.0) -> int:
    try:
        section = doc.sections[0]
        pw = section.page_width  or 0
        lm = section.left_margin or 0
        rm = section.right_margin or 0
        usable = pw - lm - rm
        if usable > 0:
            return int(usable)
    except Exception:
        pass
    return int(fallback_cm * EMU_PER_CM)


# ──────────────────────────────────────────────────────────────
#  Helper: does this paragraph contain any image?
# ──────────────────────────────────────────────────────────────

def _para_has_image(para) -> bool:
    return (
        len(para._element.findall(".//" + qn("wp:inline"))) > 0 or
        len(para._element.findall(".//" + qn("wp:anchor"))) > 0
    )


# ──────────────────────────────────────────────────────────────
#  Helper: fix line-spacing so images are not clipped
# ──────────────────────────────────────────────────────────────

def _fix_line_spacing(para) -> None:
    """Remove any fixed/exact line-height that clips images."""
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is not None:
        # Remove line / lineRule attributes; keep before/after spacing
        for attr in (qn("w:line"), qn("w:lineRule")):
            if attr in spacing.attrib:
                del spacing.attrib[attr]


# ──────────────────────────────────────────────────────────────
#  Helper: resize all inline images in a paragraph
# ──────────────────────────────────────────────────────────────

def _resize_inline_images(para, max_width_emu: int) -> None:
    for inline in para._element.findall(".//" + qn("wp:inline")):
        extent = inline.find(qn("wp:extent"))
        if extent is None:
            continue
        cx = int(extent.get("cx", 0))
        cy = int(extent.get("cy", 0))
        if cx <= 0:
            continue
        if cx > max_width_emu:
            ratio  = cy / cx
            new_cx = max_width_emu
            new_cy = int(new_cx * ratio)
            extent.set("cx", str(new_cx))
            extent.set("cy", str(new_cy))
            # Also update distT/distB/distL/distR extents in docPr if present
            docPr = inline.find(qn("wp:docPr"))
            # cNvGraphicFramePr extent
            graphic = inline.find(".//" + qn("a:graphic"))

    # Handle anchored images too
    for anchor in para._element.findall(".//" + qn("wp:anchor")):
        extent = anchor.find(qn("wp:extent"))
        if extent is None:
            continue
        cx = int(extent.get("cx", 0))
        cy = int(extent.get("cy", 0))
        if cx <= 0:
            continue
        if cx > max_width_emu:
            ratio  = cy / cx
            new_cx = max_width_emu
            new_cy = int(new_cx * ratio)
            extent.set("cx", str(new_cx))
            extent.set("cy", str(new_cy))


# ──────────────────────────────────────────────────────────────
#  Helper: insert an image into a paragraph (replaces its content)
# ──────────────────────────────────────────────────────────────

def _replace_para_with_image(
    para,
    img_bytes: bytes,
    max_width_emu: int,
    align_center: bool,
) -> None:
    """Clear paragraph text runs and insert the image."""
    # Remove all runs / hyperlinks from the paragraph (keep pPr)
    p_elem = para._element
    for child in list(p_elem):
        tag = etree.QName(child.tag).localname if child.tag != "{}" else child.tag
        if tag not in ("pPr",):
            p_elem.remove(child)

    if align_center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Determine display size
    with Image.open(io.BytesIO(img_bytes)) as img:
        img_w, img_h = img.size

    if img_w <= 0:
        return

    ratio     = img_h / img_w
    width_emu = min(max_width_emu, MAX_WIDTH_EMU)
    # python-docx add_picture accepts width in EMU via Emu()
    run = para.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=Emu(width_emu))


# ──────────────────────────────────────────────────────────────
#  Main processor class
# ──────────────────────────────────────────────────────────────

class WordImageProcessor:
    def __init__(self, doc_path: str) -> None:
        self.doc_path = doc_path
        self.doc      = Document(doc_path)

    # ── public API ────────────────────────────────────────────

    def process(
        self,
        mode:          str,
        images:        List[Tuple[str, bytes]],
        scale:         float = 0.9,
        align_center:  bool  = True,
        auto_width:    bool  = True,
        page_width_cm: float = 16.0,
    ) -> Dict[str, Any]:
        """
        Run the selected mode.
        Returns {"matched": [...], "unmatched": [...]}
        """
        # Compute target width
        if auto_width:
            page_emu = _usable_width_emu(self.doc, fallback_cm=page_width_cm)
        else:
            page_emu = int(page_width_cm * EMU_PER_CM)

        target_emu = max(1, int(page_emu * scale))

        matched:   List[str] = []
        unmatched: List[str] = []

        if mode in ("fix", "both"):
            self._fix_mode(target_emu, align_center)

        if mode in ("insert", "both"):
            matched, unmatched = self._insert_mode(images, target_emu, align_center)

        return {"matched": matched, "unmatched": unmatched}

    def save(self, output_path: str) -> None:
        self.doc.save(output_path)

    # ── private: fix mode ─────────────────────────────────────

    def _fix_mode(self, target_emu: int, align_center: bool) -> None:
        """Fix line-spacing clipping and resize all existing images."""
        for para in self._all_paragraphs():
            if _para_has_image(para):
                _fix_line_spacing(para)
                _resize_inline_images(para, target_emu)
                if align_center:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── private: insert mode ──────────────────────────────────

    def _insert_mode(
        self,
        images:       List[Tuple[str, bytes]],
        target_emu:   int,
        align_center: bool,
    ) -> Tuple[List[str], List[str]]:
        """Replace marker paragraphs with images."""
        if not images:
            return [], []

        # Build marker → bytes map  (marker = filename without extension)
        marker_map: Dict[str, Tuple[str, bytes]] = {}
        for filename, img_bytes in images:
            stem = Path(filename).stem
            marker_map[stem] = (filename, img_bytes)

        matched:   List[str] = []
        unmatched: List[str] = list(marker_map.keys())

        for para in self._all_paragraphs():
            text = para.text.strip()
            if text in marker_map:
                _, img_bytes = marker_map[text]
                try:
                    _replace_para_with_image(para, img_bytes, target_emu, align_center)
                    if text not in matched:
                        matched.append(text)
                    if text in unmatched:
                        unmatched.remove(text)
                except Exception as exc:
                    # Don't crash the whole job; keep it in unmatched
                    print(f"[processor] failed to insert image for marker '{text}': {exc}")

        return matched, unmatched

    # ── private: iterate all paragraphs (body + tables) ───────

    def _all_paragraphs(self):
        """Yield every paragraph in the document, including inside tables."""
        yield from self.doc.paragraphs
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
